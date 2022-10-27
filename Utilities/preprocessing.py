"""
This module was written for the ESF-financed project Demokratisk
Digitalisering. The project is a collaboration between
Arbetsförmedlingen (AF) and Kungliga Tekniska Högskolan (KTH). The
purpose of the project is to increase the digital competence - as
defined at https://ec.europa.eu/jrc/en/digcomp - of people working at
Arbetsförmedlingen, and by extension also job seekers. This will be
done by learning modules at KTH using OLI-Torus.

--- About this Python module ---

This module is intended to do preprocessing of the data that comes out
of OLI Torus so that it can be used by the other modules in this
repository.

Written by Alvin Gavel,
https://github.com/Alvin-Gavel/Demodigi
"""

import os
import sys

import numpy as np
import pandas as pd
import json
import matplotlib.pyplot as plt

import docx
from docx.shared import Cm
import datetime
import pytz
import xml.etree.ElementTree as et
import tqdm


plt.style.use('tableau-colorblind10')

### Internal variables of different sorts

# This is the format that I have inferred that OLI-Torus uses in the raw_analytics file
if sys.platform in ["linux", "linux2"]:
   _raw_date_format = '%B %d, %Y at %-I:%M %p UTC'
elif sys.platform == "win32":
   _raw_date_format = '%B %d, %Y at %#I:%M %p UTC'
else:
   print('Cannot figure out what OS this is. Using Unix date format and hoping for the best')
   _raw_date_format = '%B %d, %Y at %-I:%M %p UTC'
   
_xml_date_format = '%Y-%m-%d %H:%M:%S'

# For complicated reasons using the max and min functions built into
# datetime does not work with Pandas. Hence, I have picked these times
# as being safely outside the span that we are interested in.
_effective_min_date = datetime.datetime(1900, 1, 1, tzinfo = pytz.UTC)
_effective_max_date = datetime.datetime(2200, 1, 1, tzinfo = pytz.UTC)

NoneType = type(None)


class participant:
   """
   This represents a single person taking a specific learning module.
   
   Attributes
   ----------
   ID : string
   \tSome unique identifier of the participant. For privacy reasons, this
   \tis unlikely to be their actual name.
   competencies : dict of list of str
   \tA list of the competencies that the student is intended to learn, and
   \tthe individual skills that we divide them into
   answered : pandas bool DataFrame
   \tInitially empty DataFrame stating for each skill in a learning
   \tmodule whether the participant has given any answer
   answer_date : pandas datetime DataFrame
   \tInitially empty DataFrame stating for each skill in a learning
   \tmodule when the participant gave their first answer
   first_answer_date : datetime
   \tWhen the participant first gave any answer to a question
   last_answer_date : datetime
   \tWhen the participant last gave any answer to a question
   started : bool
   \tWhether the participant has answered any questions in the course module
   finished : bool
   \tWhether the participant has answered all questions in the course module
   correct_first_try : pandas bool DataFrame
   \tInitially empty DataFrame stating for each skill in a learning
   \tmodule whether the participant answered correctly on the first try.
   correct_from_start : dict of bool
   \tInitially empty dict stating for each skill whether a participant
   \tconsistently answered correctly from the first session onwards
   accumulated_by_date : dict
   \tInitially empty dict given the number of questions answered as a
   \tfunction of time.
   """
   def __init__(self, ID):
      """
      Parameters
      ----------
      ID : string
      \tDescribed under attributes
      """
      self.ID = ID
      self.competencies = {}
      self.answered = pd.DataFrame(dtype = bool)
      self.answer_date = pd.DataFrame(dtype = 'datetime64[s]')
      self.first_answer_date = None
      self.last_answer_date = None
      self.started = False
      self.finished = False
      self.correct_first_try = pd.DataFrame(dtype = bool)
      self.correct_from_start = {}
      # I would *like* to implement this one as a DataFrame, but it turns
      # out that Pandas does not accept datetime64 objects.
      self.accumulated_by_date = {}
      return

   def save_factorial_experiment_data(self, folder_path):
      """
      Save data that will be used by the python module factorial_experiment.
      """
      if folder_path[-1] != '/':
         folder_path += '/'
      f = open(folder_path + self.ID.replace('/', '_') + '.json', 'w')
      # Here we do a weird thing because json can handle Python's built-in
      # bool type but not numpy's bool_ type.
      packed = json.dumps({'ID':self.ID, 'Number of sessions':self.n_sessions(), 'Number of skills tested':self.n_skills(), 'Results':np.array(self.correct_first_try.to_numpy().tolist(), dtype=bool).tolist()})
      f.write(packed)
      f.close()
      return
      
   def save_feedback(self, save_folder_path, feedback_folder_path = 'Feedback_paragraphs/Kartläggning'):
      """
      Save a docx file of feedback that will be delivered to the participant
      by the Python module canvas_contact.
      
      At the moment, this is specific to the course module kartläggning.
      """
      def add_txt(doc, file_path):
         f = open(file_path, 'r', encoding = 'utf-8')
         contents = f.read()
         f.close()
         par = doc.add_paragraph()
         segments = contents.split('<')
         italic = False
         bold = False
         for segment in segments:
            if segment[0:2] == 'i>':
               italic = True
               segment = segment[2:]
            elif segment[0:2] == 'b>':
               bold = True
               segment = segment[2:]
            elif segment[0:3] == '/i>':
               italic = False
               segment = segment[3:]
            elif segment[0:3] == '/b>':
               bold = False
               segment = segment[3:]
               
            converted = segment.encode('Windows-1252')
            run = par.add_run(segment)
            run.italic = italic
            run.bold = bold
         return
         
      # We should never give feedback to a participant who has not actually
      # done the entire module.
      if not self.finished:
         return
         
      doc = docx.Document()
      doc.add_heading('Återkoppling på kartläggning', 0)
      add_txt(doc, '{}/Intro.txt'.format(feedback_folder_path))

      for competency_name, skills in self.competencies.items():
         n_skills = len(skills)
         n_known = 0
         for skill in skills:
            n_known += self.correct_from_start[skill]
         
         add_txt(doc, '{}/{}/Intro.txt'.format(feedback_folder_path, competency_name.replace(' ', '_')))
         add_txt(doc, '{}/{}/{}.txt'.format(feedback_folder_path, competency_name.replace(' ', '_'), n_known))
      
      add_txt(doc, '{}/Outro.txt'.format(feedback_folder_path))

      # I strongly recommend against commenting this back in. The file area
      # on Canvas is very limited, and the logo makes up about 5/6th of the
      # memory space taken up by the feedback file.
      #doc.add_picture('{}/Logo.png'.format(feedback_folder_path), width = Cm(3))

      if save_folder_path[-1] != '/':
         save_folder_path += '/'
      ID_save_folder_path = save_folder_path + self.ID.replace('/', '_')
      try:
         os.mkdir(ID_save_folder_path)
      except FileExistsError:
         pass
      doc.save('{}/Återkoppling_deltagare_{}.docx'.format(ID_save_folder_path, self.ID.replace('/', '_')))
      return
      
   def _cumulative_answers_by_date(self):
      dates_when_something_happened = {}
      for date in self.answer_date.values.flatten():
         if not pd.isnull(date):
            if not (date in dates_when_something_happened.keys()):
               dates_when_something_happened[date] = 1
            else:
               dates_when_something_happened[date] += 1
      accumulated = 0
      accumulated_by_date = {}
      for date in sorted(dates_when_something_happened.keys()):
         accumulated += dates_when_something_happened[date]
         accumulated_by_date[date] = accumulated
      self.accumulated_by_date = accumulated_by_date
      return

   def n_sessions(self):
      return self.correct_first_try.shape[0]

   def n_skills(self):
      return self.correct_first_try.shape[1]

   def plot_results_by_session(self, folder_path):
      """
      This plots the results for the participant as a function of session,
      showing both which number of skills they have answered, and for what
      number they answered correctly on the first try.
      """
      if folder_path[-1] != '/':
         folder_path += '/'
      plt.clf()
      plt.tight_layout()
      plt.plot(self.answered.index, self.answered.sum(1), label = 'Besvarade frågor')
      plt.plot(self.answered.index, self.correct_first_try.sum(1), label = 'Rätt på första försöket')
      plt.legend()
      plt.xlim(1, self.n_sessions())
      plt.ylim(0, self.n_skills())
      plt.xticks(range(1, self.n_sessions()))
      plt.savefig('{}{}_resultat_per_session.png'.format(folder_path, self.ID.replace('/', '_')))
      return

   def plot_results_by_time(self, folder_path):
      """
      This plots the number of answers given as a function of time.
      """
      # Avoid plotting if there is nothing to plot.
      if self.accumulated_by_date == {}:
         return
      
      if folder_path[-1] != '/':
         folder_path += '/'
      plt.clf()
     
      sorted_dates_accumulated = list(zip(*sorted(zip(self.accumulated_by_date.keys(), self.accumulated_by_date.values()))))
      plt.plot(sorted_dates_accumulated[0], sorted_dates_accumulated[1], 'o-', label = 'Besvarade frågor')
      plt.legend()
      plt.ylim(0, self.n_sessions() * self.n_skills())
      plt.xticks(rotation = 90)
      plt.tight_layout()
      plt.savefig('{}{}_resultat_över_tid.png'.format(folder_path, self.ID.replace('/', '_')))
      return
         
class learning_module:
   """
   This represents one learning module in the project.
  
   Attributes
   ----------
   competencies : dict of list of str
   \tA list of the competencies that the module is intended to teach, and
   \tthe individual skills that we divide them into
   skills : list of str
   \tThe skills contained in the directory of competencies
   n_skills : int
   \tThe number of skills in the learning module
   n_sessions : int
   \tThe number of sessions in the learning module. The assumption is that
   \tin each session each skill will be tested once. If nan is given, the
   \tnumber of sessions must be inferred from the OLI-Torus output.
   n_sessions_input : bool
   \tWhether a number of sessions has been supplied
   participants : dict of participant or None
   \tA list of the people taking the learning module, as identified in the
   \t'Student ID' column in the output from OLI-Torus. If None is given,
   \tthe participants must be inferred from the OLI-Torus output.
   participants_input : bool
   \tWhether a list of participants has been provided, either when 
   \tinstantiating the learning_module or afterwards from the
   n_participants : int
   \tThe number of participants
   start_date : datetime object
   \tThe time that the learning module is considered to have started.
   \tAnything before this is discarded.
   end_date : datetime object
   \tThe time that the learning module is considered to have ended.
   \tAnything after this is discarded.
   section_slug : str
   \tIdentifier internal to the raw_analytics file that describes the version
   \tof the learning module. If this is specified, all entries that do not
   \tmatch are discarded.
   flags : pandas DataFrame
   \tFlags that note whether each participant has:
   \t1. Started, by answering at least one question
   \t2. Finished the learning module   
   raw_analytics : pandas dataframe
   \tThe exact contents of the raw_analytics file that results are read from.
   \tThis is initially empty.      
   full_results : pandas DataFrame
   \tThe results from the learning module as output by OLI-Torus
   results_read : bool
   \tWhether results have been read from a file
   accumulated_by_date : dict
   \tInitially empty dict given the number of questions answered as a
   \tfunction of time.
   """
   def __init__(self, competencies, n_sessions = np.nan, participants = None, start_date = _effective_min_date, end_date = _effective_max_date, section_slug = None):
      """
      Parameters
      ----------
      competencies : dict of lists of str
      \tDescribed under attributes
      
      Optional parameters
      -------------------
      n_sessions : int
      \tDescribed under attributes
      participants : dict of participant or None
      \tDescribed under attributes
      start_date : DateTime or None
      \tIf specified, gives the time at which the module is taken to have
      \tstarted. All data from before this is ignored when importing results
      end_date : DateTime or None
      \tIf specified, gives the time at which the module is taken to have
      \tended. All data from after this is ignored when importing results
      """
      self.competencies = competencies
      self.skills = []
      for competence, skills in competencies.items():
         self.skills += skills
      self.n_skills = len(self.skills)
      self.n_sessions = n_sessions
      self.n_sessions_input = np.isfinite(self.n_sessions)
      self.participants = participants
      if self.participants == None:
         self.participants_input = False
         self.n_participants = np.nan
      else:
         self.participants_input = True
         self.n_participants = len(participants)
      self.start_date = start_date
      self.end_date = end_date
      self.section_slug = section_slug
      self.flags = pd.DataFrame()
      self.raw_analytics = None
      self.full_results = None
      self.results_read = False
      self.accumulated_by_date = {}
      return

   ### Functions for handling data regarding individual participants

   def read_participant_IDs(self, filepath):
      """
      Reads a list of the participant IDs. The file is assumed to consist of
      a single column of IDs.
      """
      file_ending = filepath.split('.')[-1]
      if file_ending == 'txt':
         f = open(filepath)
         IDs = [word.strip() for word in f]
         f.close()
      elif file_ending == 'xlsx':
         dataframe = pd.read_excel(filepath, header = 0, dtype = str)
         IDs = list(dataframe['user_id'])
      else:
         print('Cannot recognise file type of {}'.format(filepath))
      self.participants = {}
      for ID in IDs:
         self.participants[ID] = participant(ID) 
      self.participants_input = True
      self.n_participants = len(self.participants)
      return

   def import_raw_analytics(self, filepath):
      """
      Import the raw_analytics.tsv file given by OLI Torus and pick out the 
      relevant information.
      """
      raw = pd.read_csv(filepath, sep='\t')
      
      raw['Date Created']= pd.to_datetime(raw['Date Created'])
      cleaned = raw.astype({'Student ID': str}) # This sometimes gets interpreted as int
      if self.section_slug != None:
         cleaned = cleaned[cleaned['Section Slug'] == self.section_slug]
      
      self.raw_data_full = pd.DataFrame(data={'Student ID':cleaned['Student ID'], 'Date Created':cleaned['Date Created'], 'Activity Title': cleaned['Activity Title'], 'Attempt Number': cleaned['Attempt Number'], 'Correct?': cleaned['Correct?']})
      self.raw_data = self.raw_data_full[(self.start_date < self.raw_data_full['Date Created']) & (self.raw_data_full['Date Created'] < self.end_date)].reset_index()
      self.results_read = True
      return
      
   def import_datashop(self, filepaths, verbose = False):
      """
      This imports an XML file following the format specified at
      
      https://pslcdatashop.web.cmu.edu/dtd/guide/tutor_message_dtd_guide_v4.pdf
      
      The assumption is that for each problem there will be a context_message
      followed by one or more pairs of first tool_message and tutor_message.
      Internally, these are ordered in time, but the whole batch of context,
      tool and tutor messages need not be ordered with respect to others.
      
      The import function can be given one or many filepaths. In the latter
      case, they must come from courses that are considered separate in OLI
      Torus, but which contain questions with the same names.
      """
      def match_skill(problem_name):
         """
         See if the problem name matches the format of a question on any of the
         skills we want to test.
         """
         matched_skill = ''
         for skill in self.skills:
            for session in range(self.n_sessions):
               mixedcase_name = '{}_Q{}'.format(skill, session + 1)
               if problem_name == mixedcase_name.lower():
                  matched_skill = mixedcase_name
         return matched_skill

      if type(filepaths) == str:
         filepaths = [filepaths]

      # We make a first run to pick out the relevant information and store it
      # in a way that will allow us to pick out the attempt number.
      answers = {}
      for filepath in filepaths:
         tree = et.parse(filepath)
         root = tree.getroot()
      
         # Some questions do not match the standard format for problems in the
         # XML file. They should simply be ignored.
         wait_for_next_context_message = False
      
         for child in root:
            if child.tag == 'context_message':
               meta = child.findall('meta')[0]
               anon_id = meta.findall('user_id')[0].text
               time_str = meta.findall('time')[0].text
               # The Datashop file says that the timezone is GMT, but that is the same thing as UTC
               time = datetime.datetime.strptime(time_str, _xml_date_format).replace(tzinfo=datetime.timezone.utc)
            
               try:
                  full_problem_name = child.findall('dataset')[0].findall('level')[0].findall('level')[0].findall('problem')[0].findall('name')[0].text
               except IndexError:
                  wait_for_next_context_message = True
                  continue
               wait_for_next_context_message = False
            
               problem_name = full_problem_name.split(' ')[1][:-1]

               if not (anon_id in answers.keys()):
                  answers[anon_id] = {}
               if not (time in answers[anon_id].keys()):
                  answers[anon_id][time] = {}
               if not (problem_name in answers[anon_id][time].keys()):
                  answers[anon_id][time][problem_name] = []
            elif child.tag == 'tool_message' and (not wait_for_next_context_message):
               pass
            elif child.tag == 'tutor_message' and (not wait_for_next_context_message):
               action_eval = child.findall('action_evaluation')[0].text
            
               if action_eval == 'CORRECT':
                  is_correct = True
               elif action_eval == 'INCORRECT':
                  is_correct = False
               else:
                  print('Cannot figure out if action was completed correctly or not')
               answers[anon_id][time][problem_name].append(is_correct)

      # Prepare to store the information in better structured ways
      IDs = []
      answer_dates = []
      next_dates = []
      correct = []
      activity_titles = []
      
      batches_full = {}
      for anon_id, times in answers.items():
         times_problem_names = sorted(times.items())
         n_times_problem_names = len(times_problem_names)
         
         batches_full[anon_id] = {}
         for i in range(n_times_problem_names):
            time, problem_names = times_problem_names[i]
            if i < n_times_problem_names-1:
               next_time, dummy = times_problem_names[i+1]
            else:
               next_time == _effective_max_date
            
            batches_full[anon_id][(time, next_time)] = {}
            for problem_name, answers in problem_names.items():
               matched_skill = match_skill(problem_name)
               if matched_skill == '':
                  continue
                  
               batches_full[anon_id][(time, next_time)][matched_skill] = []               
               for is_correct in answers:
                  IDs.append(anon_id)
                  activity_titles.append(matched_skill)
                  correct.append(is_correct)
                  answer_dates.append(time)
                  next_dates.append(next_time)
         
                  batches_full[anon_id][(time, next_time)][matched_skill].append(is_correct)
      
            if batches_full[anon_id][(time, next_time)] == {}:
               batches_full[anon_id].pop((time, next_time), None)
      
      batches = {}
      for anon_id, timespans in batches_full.items():
         batches[anon_id] = {}
         for timespan in timespans.keys():
            if self.start_date < timespan[0] < self.end_date:
               batches[anon_id][timespan] = batches_full[anon_id][timespan]
         if batches[anon_id] == {}:
            batches.pop(anon_id, None)
      
      self.xml_dict_full = batches_full
      self.xml_dict = batches
      # Note that the next date might make use of feedback and other things which are not stored in the dataframe itself.
      self.xml_dataframe_full = pd.DataFrame(data={'Student ID (lowercase)':IDs, 'Date Created':answer_dates, 'Next Date':next_dates, 'Activity Title': activity_titles, 'Correct?': correct})
      self.xml_dataframe = self.xml_dataframe_full[(self.start_date < self.xml_dataframe_full['Date Created']) & (self.xml_dataframe_full['Date Created'] < self.end_date)].reset_index()
      self.results_read = True
      return

   def _infer_mapping_pseudonym_ID(self, verbose = True):
      """
      See the internal DD document 'Datashop och raw analytics' for an
      explanation of what this method does and why. Note that at the moment
      *IT DOES NOT WORK*, so it will need fixing if you want to use it.
      """
      n_lowercaseID = len(self.xml_dict)
      
      pseudonyms = list(set(self.raw_data['Student ID']))
      n_pseudonym = len(pseudonyms)

      mapping = {'Student ID (lowercase)':[],'Pseudonym':[],'Match reliable':[],'Match percentage':[], 'n matches':[], 'Total occurrences':[], 'Second best match percentage':[]}      

      mapping_pseudonym_lowercaseID = {}
      mapping_lowercaseID_pseudonym = {}
      
      unmapped = {'Student ID (lowercase)':[], 'Match percentage':[]}
   
      if verbose:
         print('Figuring out mapping between pseudonyms in raw_analytics file and IDs in Datashop file.')
         print('Comparing {} pseudonyms to {} IDs'.format(n_pseudonym, n_lowercaseID))
         print('This may take a while...')
         
      all_matches = 0
      reliable_matches = 0
      for lowercaseID, batches in tqdm.tqdm(self.xml_dict.items()):
      
         n_matches = []
         total_occurrences = []
         match_percentages = []
         for pseudonym in pseudonyms:
            pseudonym_entries = self.raw_data[self.raw_data['Student ID'] == pseudonym]
         
            n_problems = 0
            n_problems_matched = 0
            for time_span, problem_names in batches.items():
               not_too_early = time_span[0] <= pseudonym_entries['Date Created'] - datetime.timedelta(seconds=60)
               not_too_late = pseudonym_entries['Date Created'] <= time_span[1]
               
               for problem_name, batch_answers in problem_names.items():

                  same_problem = pseudonym_entries['Activity Title'] == problem_name
                  pseudonym_answers = pseudonym_entries[not_too_early & not_too_late & same_problem]['Correct?'].values
                
                  as_many_answers = len(pseudonym_answers) == len(batch_answers)
                  as_many_correct = sum(pseudonym_answers) == sum(batch_answers)

                  matched = as_many_answers and as_many_correct
                  
                  n_problems_matched += matched
                  n_problems += 1
            n_matches.append(n_problems_matched)
            total_occurrences.append(n_problems)
            match_percentages.append(n_problems_matched / n_problems)
        
         match_percentages = np.asarray(match_percentages)
            
         best_match_index = np.argmax(match_percentages)
         best_match_percentage = match_percentages[best_match_index]
         matched_pseudonym = pseudonyms[best_match_index]
         best_n_matches = n_matches[best_match_index]
         best_total_occurrences = total_occurrences[best_match_index]
         
         match_percentages[best_match_index] = -np.inf
         second_best_match_index = np.argmax(match_percentages)
         second_best_match_percentage = match_percentages[second_best_match_index]

         if best_match_percentage >= 0.5:
            all_matches += 1
           
            reliable = best_match_percentage - second_best_match_percentage >= 0.5 and best_n_matches > 5
            reliable_matches += reliable
         
            pseudonyms.remove(matched_pseudonym)
            mapping['Student ID (lowercase)'].append(lowercaseID.replace('@arbetsformedlingen.se', ''))
            mapping['Pseudonym'].append(matched_pseudonym)
            mapping['Match reliable'].append(reliable)      
            mapping['Match percentage'].append(best_match_percentage)
            mapping['n matches'].append(best_n_matches)
            mapping['Total occurrences'].append(best_total_occurrences)
            mapping['Second best match percentage'].append(second_best_match_percentage)
            if reliable:
               mapping_pseudonym_lowercaseID[matched_pseudonym] = lowercaseID.replace('@arbetsformedlingen.se', '')
               mapping_lowercaseID_pseudonym[lowercaseID.replace('@arbetsformedlingen.se', '')] = matched_pseudonym
         else:
            unmapped['Student ID (lowercase)'].append(lowercaseID.replace('@arbetsformedlingen.se', ''))
            unmapped['Match percentage'].append(best_match_percentage)
            
         if len(pseudonyms) == 0:
            break
      self.full_mapping = pd.DataFrame(data=mapping)
      self.mapping = self.full_mapping[self.full_mapping['Match reliable']]
      self.mapping_pseudonym_lowercaseID = mapping_pseudonym_lowercaseID
      self.mapping_lowercaseID_pseudonym = mapping_lowercaseID_pseudonym
      
      self.unmapped = pd.DataFrame(data=unmapped)
      
      self.unmatched_pseudonyms = []
      for pseudonym in pseudonyms:
         if not (pseudonym in self.mapping['Pseudonym'].values):
            self.unmatched_pseudonyms.append(pseudonym)

      self.unmatched_lowercaseIDs = []
      for lowercaseID in self.xml_dict.keys():
         if not (lowercaseID.replace('@arbetsformedlingen.se', '') in self.mapping['Student ID (lowercase)'].values):
            self.unmatched_lowercaseIDs.append(lowercaseID)
            
      if verbose:
         print('There are {} unique pseudonyms in the raw_analytics file'.format(n_pseudonym))
         print('Of these, {} could be matched at all to IDs in the Datashop file'.format(all_matches))
         print('Of these, {} could be reliably matched to IDs in the Datashop file'.format(reliable_matches))
         print('There are {} unique IDs in the Datashop file'.format(n_lowercaseID))
         print('Of these, {} could be reliably matched to pseudonyms in the raw_analytics file'.format(len(mapping_lowercaseID_pseudonym)))
         print('Another   {} could not'.format(len(self.unmatched_lowercaseIDs)))
      return


   def import_data(self, raw_analytics_path, xml_path, verbose = False, previous_mapping_path = None):
      """
      Import data from the raw_analytics and Datashop files output by OLI
      Torus and extract from them the data we need. Neither file individually
      contains all the data, but together they do.
      """
      self.import_raw_analytics(raw_analytics_path)
      self.import_datashop(xml_path)
      self._infer_mapping_pseudonym_ID(verbose = verbose)#, previous_mapping_path = previous_mapping_path)
      
      student_ids = []
      date_created = []
      activity_title = []
      attempt_number = []
      correct = []
      for i in range(self.raw_data.shape[0]):
         pseudonym = self.raw_data['Student ID'][i]
         if pseudonym in self.mapping_pseudonym_lowercaseID.keys():
            student_ids.append(self.mapping_pseudonym_lowercaseID[pseudonym])
            date_created.append(self.raw_data['Date Created'][i])
            activity_title.append(self.raw_data['Activity Title'][i])
            attempt_number.append(self.raw_data['Attempt Number'][i])
            correct.append(self.raw_data['Correct?'][i])

      self.full_results = pd.DataFrame(data={'Student ID':student_ids, 'Date Created':date_created, 'Activity Title':activity_title,'Attempt Number':attempt_number, 'Correct?':correct})
      return


   def _setup_result_reading(self, participant):
      """
      Create the empty dataframes and dummy dates necessary before filling in
      results for the participant.
      """
      if not self.n_sessions_input:
         print('Inferring number of sessions from OLI-Torus output')
         self.infer_n_sessions_from_full_results()

      participant.competencies = self.competencies #Throwing this piece of data around like this is probably bad. Might fix later
      participant.answered = pd.DataFrame(columns = self.skills, index = range(1, self.n_sessions + 1), dtype = bool)
      participant.answer_date = pd.DataFrame(columns = self.skills, index = range(1, self.n_sessions + 1), dtype = 'datetime64[s]')
      participant.correct_first_try = pd.DataFrame(columns = self.skills, index = range(1, self.n_sessions + 1), dtype = bool)
      # Using the max and min of datetime does not work together with Pandas
      participant.first_answer_date = _effective_max_date
      participant.last_answer_date = _effective_min_date
      return
      
   def _read_participant_results_from_combined(self, participant):
      """
      Find out, for each question, whether a specific participant got it
      right on the first try.
      """
      self._setup_result_reading(participant)
      correct_participant = self.full_results[self.full_results['Student ID'] == participant.ID.lower()]
      n_answers = 0
      for skill in self.skills:
         for session in range(1, self.n_sessions + 1):
            try:
               correct_skill = correct_participant[correct_participant['Activity Title'] == '{}_Q{}'.format(skill, session)]
               first_try_index = correct_skill['Attempt Number'] == 1
               # If nothing was found, this will throw an IndexError
               first_try_date = correct_skill['Date Created'][first_try_index].to_numpy()[0]
               has_answered = True
               correct = correct_skill['Correct?'][first_try_index].to_numpy()[0]
               n_answers += 1
               
               if first_try_date < participant.first_answer_date:
                  participant.first_answer_date = first_try_date
               if first_try_date > participant.last_answer_date:
                  participant.last_answer_date = first_try_date
            except IndexError:
               has_answered = False
               correct = False
               first_try_date = None
            participant.answered.loc[session, skill] = has_answered
            participant.correct_first_try.loc[session, skill] = correct
            participant.answer_date.loc[session, skill] = first_try_date
         participant.correct_from_start[skill] = np.all(participant.correct_first_try.loc[:, skill])
      participant._cumulative_answers_by_date()
      participant.started = n_answers > 0
      participant.finished = n_answers == self.n_sessions * self.n_skills
      self.flags.loc[participant.ID, 'started'] = participant.started
      self.flags.loc[participant.ID, 'finished'] = participant.finished
      return
      
   def _read_participant_results_from_xml(self, participant):
      """
      Try to find out, for each question, whether a specific participant got
      it right on the first try.
      
      Since the Datashop file does not contain timestamps for individual
      answer attempts, this function will assume that a participant got it
      right on the first try iff they gave one answer, which was correct. In
      those cases where a participant answered correctly on the first try and
      then checked the other alternatives out of curiosity, they will be
      flagged as having answered incorrectly. This is *very bad* but we may
      not have any viable alternative.
      """
      
      self._setup_result_reading(participant)
      correct_participant = self.xml_dataframe[self.xml_dataframe['Student ID (lowercase)'] == participant.ID.lower() + '@arbetsformedlingen.se']
      n_answers = 0
      for skill in self.skills:
         for session in range(1, self.n_sessions + 1):
            correct_skill = correct_participant[correct_participant['Activity Title'] == '{}_Q{}'.format(skill, session)]
               
            if len(correct_skill) == 0:
               has_answered = False
               correct = False
               first_try_date = None
            else:            
               has_answered = True
               n_answers += 1
            
               first_try_date = _effective_max_date
               for date in correct_skill['Date Created']:
                  if first_try_date > date:
                     first_try_date = date
               
               correct = np.all(correct_skill[correct_skill['Date Created'] == first_try_date]['Correct?'])
               
               if first_try_date < participant.first_answer_date:
                  participant.first_answer_date = first_try_date
               if first_try_date > participant.last_answer_date:
                  participant.last_answer_date = first_try_date

            participant.answered.loc[session, skill] = has_answered
            participant.correct_first_try.loc[session, skill] = correct
            participant.answer_date.loc[session, skill] = first_try_date
         participant.correct_from_start[skill] = np.all(participant.correct_first_try.loc[:, skill])
      participant._cumulative_answers_by_date()
      participant.started = n_answers > 0
      participant.finished = n_answers == self.n_sessions * self.n_skills
      self.flags.loc[participant.ID, 'started'] = participant.started
      self.flags.loc[participant.ID, 'finished'] = participant.finished
      return
      
   def read_participants_results(self, database = 'combined', verbose = True):
      """
      Find out, for each question, whether the participants got it right on
      the first try.
      """
      # If you simply test '[...] == None' Pandas will complain that
      # dataframes have ambiguous equality.
      if (database == 'combined' and type(self.full_results) == NoneType) or (database == 'datashop' and type(self.xml_dataframe) == NoneType):
         print('No results have been read!')
      else:
         if verbose:
            print("Reading participants' results from {} database. This may take a while...".format(database))
         for participant in tqdm.tqdm(self.participants.values()):
            if database == 'combined':
               self._read_participant_results_from_combined(participant)
            elif database == 'datashop':
               self._read_participant_results_from_xml(participant)
            else:
               print('Cannot recognise database {}'.format(database))
               return
         self.accumulated_by_date = self._cumulative_answers_by_date(self.participants.values())
      return

   def export_full_results(self, file_path):
      """
      Export the full results dataframe as a csv file.
      
      This is mostly intended to make visual inspection easier.
      """
      if type(self.full_results) == NoneType:
         if type(self.xml_dataframe) == NoneType:
            print('No results have been read!')
         else:
            print('Full results have not been read!\nThese are necessary to give meaningful individual results.')
      else:
         self.full_results.to_csv(file_path, index = False)
      return

   def export_individual_results(self, folder_path, verbose = True):
      """
      Export the results for each individual participant in a format which is
      legible to the factorial_experiment module.
      """
      if not self.results_read:
         print('There are no results to save!')
      else:
         if verbose:
            print("Exporting participants' results. This may take a while...")
         for participant in tqdm.tqdm(self.participants.values()):
            participant.save_factorial_experiment_data(folder_path)
      return

   def export_individual_feedback(self, folder_path, verbose = True):
      """
      Export the feedback for each individual participant, which can then be
      uploaded by the feedback module to Canvas.
      """
      if folder_path[-1] != '/':
         folder_path += '/'
      
      if not self.results_read:
         print('There are no results to give feedback on!')
      else:
         if verbose:
            print("Exporting participants' feedback. This may take a while...")
         for participant in tqdm.tqdm(self.participants.values()):
            participant.save_feedback(folder_path)
      return

   def export_IDs(self, file_path):
      """
      Write a file of participant IDs, which can be read by the
      factorial_experiment module, or by this module.
      
      There should not be any need to use this function, except when
      participants have been inferred from the results file.
      """
      f = open(file_path, 'w')
      # Here we do a weird thing because json can handle Python's built-in
      # bool type but not numpy's bool_ type.
      packed = json.dumps({'IDs':list(self.participants.keys())})
      f.write(packed)
      f.close()
      return      

   def export_datashop(self, file_path):
      """
      Export the data interpreted from the XML file
      """
      self.xml_dataframe.to_csv(file_path, index = False)
      return

   def export_mapping(self, file_path):
      """
      Export the mapping between IDs as given in the Datashop and raw_analytics
      files
      """
      self.mapping.to_csv(file_path, index = False)
      return
      
   def export_full_mapping(self, file_path):
      """
      Export the mapping between IDs as given in the Datashop and raw_analytics
      files, including tentative matches
      """
      self.full_mapping.to_csv(file_path, index = False)
      return
      
   def export_unmapped(self, file_path):
      """
      Export a dataframe with those IDs that could not be mapped
      """
      self.unmapped.to_csv(file_path, index = False)
      return
      
   def export_SCB_data(self, file_path):
      """
      Write a preliminary csv file to be delivered to Statistiska
      Centralbyrån (SCB), which contains the columns:
      
      user name, estimated time, starting date, finishing date
      
      The estimated time is always given as 30 minutes. Note that this file
      is only preliminary. It will need to go through a second stage of
      processing where the user names are replaced with the personal identity
      numbers of the participants.
      """
      sorted_IDs = sorted(self.participants.keys())
      first_answer_dates = []
      last_answer_dates = []
      for ID in sorted_IDs:
         participant = self.participants[ID]
         if participant.started:
            first_answer_dates.append(participant.first_answer_date.date())
         else:
            first_answer_dates.append('Ej börjat')
         if participant.finished:
            last_answer_dates.append(participant.last_answer_date.date())
         else:
            last_answer_dates.append('Ej klar')

      SCB_data = pd.DataFrame(columns = ['Användarnamn', 'Uppskattad tid', 'Startdatum', 'Avslutsdatum'])
      SCB_data['Startdatum'] = first_answer_dates
      SCB_data['Avslutsdatum'] = last_answer_dates
      SCB_data['Användarnamn'] = sorted_IDs
      SCB_data['Uppskattad tid'] = '30 min'
      SCB_data.to_csv(file_path, index = False)
      return

   ### Functions for handling data regarding groups of participants
   
   def _cumulative_answers_by_date(self, participants):
      dates_when_something_happened = {}
      for participant in participants:
         for date in participant.answer_date.values.flatten():
            if not pd.isnull(date):
               if not (date in dates_when_something_happened.keys()):
                  dates_when_something_happened[date] = 1
               else:
                  dates_when_something_happened[date] += 1
      accumulated = 0
      accumulated_by_date = {}
      for date in sorted(dates_when_something_happened.keys()):
         accumulated += dates_when_something_happened[date]
         accumulated_by_date[date] = accumulated
      return accumulated_by_date
   
   ### Functions for inspecting data

   def describe_module(self):
      """
      Output some basic numerical data about how the participants have done
      as a group.
      """
      if not self.participants_input:
         print('No participants have been read!')
      else:
         print('There are {} participants:'.format(len(self.participants)))
         n_started = 0
         n_finished = 0
         for ID, participant in sorted(self.participants.items()):
            if self.results_read:
               if participant.finished:
                  n_started += 1
                  n_finished += 1
               elif participant.started:
                  n_started += 1
         print('{} have started'.format(n_started))
         print('{} have finished'.format(n_finished))
      return

   def describe_participants(self):
      """
      Give a short summary of who the participants are and how far they have
      gotten in the course.
      """
      if not self.participants_input:
         print('No participants have been read!')
      else:
         print('There are {} participants:'.format(len(self.participants)))
         for ID, participant in sorted(self.participants.items()):
            if self.results_read:
               if participant.finished:
                  status_string = 'Has finished module'
               elif participant.started:
                  status_string = 'Has started work'
               else:
                  status_string = 'Has not started'
            else:
               status_string = 'No results known'
            print('   {}: {}'.format(ID, status_string))
      return
      
   def plot_individual_results(self, folder_path):
      """
      Plot the results for each individual participant.
      """
      for participant in self.participants.values():
         participant.plot_results_by_session(folder_path)
         participant.plot_results_by_time(folder_path)
      return

   def plot_results_by_time(self, folder_path):
      """
      This plots the number of answers given as a function of time.
      """
      if folder_path[-1] != '/':
         folder_path += '/'
      plt.clf()
     
      sorted_dates_accumulated = list(zip(*sorted(zip(self.accumulated_by_date.keys(), self.accumulated_by_date.values()))))
      if len(sorted_dates_accumulated) == 0:
         print('Nothing to plot')
         return
      plt.plot(sorted_dates_accumulated[0], sorted_dates_accumulated[1], '-', label = 'Besvarade frågor')
      plt.legend()
      plt.ylim(0, self.n_sessions * self.n_skills * self.n_participants)
      plt.xticks(rotation = 90)
      plt.tight_layout()
      plt.savefig('{}Resultat_över_tid.png'.format(folder_path))
      return

   def plot_performance_by_n_correct_for_competence(self, folder_path, competence, threshold = 4/6):
      """
      Make a barplot showing how many participants got 0, 1, 2... answers
      right on the first try of the first session and onwards, for a given
      competence
      """
      if folder_path[-1] != '/':
         folder_path += '/'

      n_skill_for_this_competency = len(self.competencies[competence])

      x = []
      already_known = []
      colors = []
      for i in range(n_skill_for_this_competency + 1):
         x.append(i)
         already_known.append(0)
         
         if i / n_skill_for_this_competency <= threshold:
            colors.append('orange')
         else:
            colors.append('green')
      
      for participant in self.participants.values():
         n_correct = 0
         for skill in self.competencies[competence]:
            n_correct += participant.correct_from_start[skill]
         already_known[n_correct] += 1
         
      for extension in ['', '_normaliserad']:
         plt.clf()
         plt.bar(x, already_known, color = colors, edgecolor='black')
         if extension == '_normaliserad':
            plt.ylim(0, self.n_participants)
         else:
            plt.ylim(0, max(already_known))
         plt.xlabel('Antal rätt')
         plt.ylabel('Antal deltagare')
         plt.title(competence)
         plt.tight_layout()
         plt.savefig('{}{}_per_antal_rätt{}.png'.format(folder_path, competence.replace(' ', '_'), extension))
      return

   def plot_performance_per_skill_for_competence(self, folder_path, competence):
      """
      Make a barplot showing how many participants got a given skill right on
      the first try, for a given competence
      """
      if folder_path[-1] != '/':
         folder_path += '/'

      n_skill_for_this_competency = len(self.competencies[competence])

      x = []
      already_known = []
      for i in range(n_skill_for_this_competency):
         x.append(i)
         already_known.append(0)
      
      for participant in self.participants.values():
         for i in range(n_skill_for_this_competency):
            skill = self.competencies[competence][i]
            already_known[i] += np.all(participant.correct_first_try.loc[:, skill])
            
      for extension in ['', '_normaliserad']:
         plt.clf()
         plt.bar(x, already_known)
         plt.xticks(ticks=x, labels=self.competencies[competence], rotation=90)
         if extension == '_normaliserad':
            plt.ylim(0, self.n_participants)
         else:
            plt.ylim(0, max(already_known))
         plt.ylabel('Deltagare med rätt på första försöket')
         plt.title(competence)
         plt.tight_layout()
         plt.savefig('{}{}_per_skill{}.png'.format(folder_path, competence.replace(' ', '_'), extension))
      return

   def plot_initial_performance(self, folder_path):
      """
      This makes bar plots showing the performance of the participants
      """
      for competence in self.competencies.keys():
         self.plot_performance_by_n_correct_for_competence(folder_path, competence)
         self.plot_performance_per_skill_for_competence(folder_path, competence)
      return

   def plot_results(self, folder_path):
      """
      This plots everything that can be plotted
      """
      self.plot_individual_results(folder_path)
      self.plot_results_by_time(folder_path)
      self.plot_initial_performance(folder_path)
      return
   
   ### Function for inferring parameters when they are not available
   ### (as a rule, you should not need to use these unless something has
   ### gone wrong)
   
   def infer_participants_from_full_results(self):
      """
      Figure out who the participants are from a file of results.
      
      NOTE: You should never actually need to use this
      """
      if not self.results_read:
         print('Must read a file of results first!')
      else:
         inferred_participant_IDs = set(self.full_results['Student ID'])
         self.participants = {}
         for ID in inferred_participant_IDs:
            self.participants[ID] = participant(ID)
      self.n_participants = len(self.participants)
      self.participants_input = True
      return
      
   def infer_n_sessions_from_full_results(self, verbose = False):
      """
      Figure out how many sessions there are from a file of results.
      
      NOTE: You should never actually need to use this
      """
      if not self.results_read:
         print('Must read a file of results first!')
      else:
         n_sessions = {}
         for skill in self.skills:
            n_sessions[skill] = 0
            activities = self.full_results['Activity Title'].to_numpy()
            while '{}_Q{}'.format(skill, n_sessions[skill] + 1) in activities:
               n_sessions[skill] += 1
      if verbose:
         print('Number of sessions registered:')
         for skill_name, n in n_sessions.items():
            print('{}: {}'.format(skill_name, n))
               
      self.n_sessions = max(n_sessions.values())
      self.n_sessions_input = True
      return
